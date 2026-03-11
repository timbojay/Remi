"""Export biography to various formats: Markdown, HTML, TXT, PDF, DOCX."""

import os
import tempfile
from datetime import date


async def export_biography(content: str, format: str, user_name: str = "User") -> tuple[bytes, str, str]:
    """Export biography content to the specified format.

    Returns: (file_bytes, filename, content_type)
    """
    title = f"The Life of {user_name}"
    today = date.today().isoformat()

    if format == "markdown" or format == "md":
        text = f"# {title}\n\n*Generated on {today}*\n\n{content}"
        filename = f"biography_{user_name.lower().replace(' ', '_')}.md"
        return text.encode("utf-8"), filename, "text/markdown"

    elif format == "txt":
        # Strip markdown formatting
        import re
        plain = re.sub(r'#{1,6}\s+', '', content)
        plain = re.sub(r'\*\*(.*?)\*\*', r'\1', plain)
        plain = re.sub(r'\*(.*?)\*', r'\1', plain)
        text = f"{title}\nGenerated on {today}\n{'=' * 40}\n\n{plain}"
        filename = f"biography_{user_name.lower().replace(' ', '_')}.txt"
        return text.encode("utf-8"), filename, "text/plain"

    elif format == "html":
        import re
        # Simple markdown to HTML conversion
        html_content = content
        html_content = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_content)
        html_content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html_content)
        # Wrap paragraphs
        paragraphs = html_content.split('\n\n')
        html_content = '\n'.join(
            f'<p>{p.strip()}</p>' if not p.strip().startswith('<h') else p.strip()
            for p in paragraphs if p.strip()
        )

        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 700px; margin: 40px auto; padding: 0 20px; line-height: 1.7; color: #333; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #2c3e50; margin-top: 30px; }}
        p {{ text-align: justify; }}
        .date {{ color: #7f8c8d; font-style: italic; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="date">Generated on {today}</p>
    {html_content}
</body>
</html>"""
        filename = f"biography_{user_name.lower().replace(' ', '_')}.html"
        return html.encode("utf-8"), filename, "text/html"

    elif format == "pdf":
        # Generate HTML first, then convert
        html_bytes, _, _ = await export_biography(content, "html", user_name)
        # PDF requires additional dependencies - for now return HTML with note
        filename = f"biography_{user_name.lower().replace(' ', '_')}.pdf"
        # TODO: Use weasyprint or similar for PDF generation
        return html_bytes, filename.replace('.pdf', '.html'), "text/html"

    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re

            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(f"Generated on {today}", style='Subtitle')

            # Parse markdown into paragraphs and headings
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    # Strip markdown formatting
                    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                    clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                    doc.add_paragraph(clean)

            # Save to bytes
            import io
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            filename = f"biography_{user_name.lower().replace(' ', '_')}.docx"
            return buffer.read(), filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except ImportError:
            # Fallback to markdown if python-docx not installed
            return await export_biography(content, "markdown", user_name)

    else:
        raise ValueError(f"Unsupported format: {format}")
